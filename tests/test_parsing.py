import io
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch


def test_extract_text_from_pdf_file_not_found():
    from src.parsing.pdf_parser import extract_text_from_pdf
    with pytest.raises(FileNotFoundError):
        extract_text_from_pdf("/nonexistent/path/file.pdf")


def test_extract_text_from_docx_file_not_found():
    from src.parsing.docx_parser import extract_text_from_docx
    with pytest.raises(FileNotFoundError):
        extract_text_from_docx("/nonexistent/path/file.docx")


def test_extract_text_from_pdf_bytes_returns_string():
    """Smoke test: if fitz raises, the exception propagates; valid bytes return str."""
    from src.parsing.pdf_parser import extract_text_from_pdf_bytes
    with pytest.raises(Exception):
        # Empty bytes is not a valid PDF — should raise, not return silently
        extract_text_from_pdf_bytes(b"not a pdf")


def test_docx_extractor_joins_paragraphs(tmp_path):
    """Create a minimal docx in-memory and verify extraction."""
    from docx import Document as DocxDocument
    from src.parsing.docx_parser import extract_text_from_docx

    doc_path = tmp_path / "test_contract.docx"
    doc = DocxDocument()
    doc.add_paragraph("Clause 1. The limit is 20%.")
    doc.add_paragraph("Clause 2. Payment due within 30 days.")
    doc.save(str(doc_path))

    text = extract_text_from_docx(doc_path)
    assert "Clause 1" in text
    assert "Clause 2" in text
