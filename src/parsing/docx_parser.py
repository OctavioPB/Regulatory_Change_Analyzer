import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract plain text from a .docx file.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Paragraphs joined with newlines, tables represented as tab-separated rows.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error("Failed to open %s: %s", path, exc)
        raise

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    text = "\n".join(parts)
    logger.debug("Extracted %d chars from %s", len(text), path.name)
    return text
