"""Automated addendum drafting — LLM-powered legal document generation."""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import AsyncIterator

import anthropic
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from src.config import settings

logger = logging.getLogger(__name__)

# ── Prompt ────────────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """You are a specialized legal document drafting AI for a financial institution's compliance department.

Your task: Draft a FORMAL CONTRACT ADDENDUM in response to the regulatory change(s) provided.

FORMAT REQUIREMENTS:
1. HEADER BLOCK — document title, effective date placeholder, parties placeholder
2. RECITALS — WHEREAS clauses citing each regulatory trigger with source, date, and article
3. NOW THEREFORE — operative transition phrase
4. ARTICLES — one per clause modification, each containing:
   - "ARTICLE N — AMENDMENT TO [CLAUSE REF]" as section heading
   - The specific amendment instruction using SHALL/IS HEREBY language
   - The new clause text verbatim or as a diff ("...is deleted and replaced with: ...")
5. GENERAL PROVISIONS article containing:
   - Remaining provisions in full force and effect
   - Governing law (same as underlying contract, or [JURISDICTION] if unknown)
   - Counterparts clause
6. EXECUTION / SIGNATURE BLOCK with lines for two authorized signatories

STYLE:
- Formal legal register: SHALL, HEREIN, WHEREAS, NOW THEREFORE, IN WITNESS WHEREOF
- Do not use bullet points inside operative clauses — use numbered sub-articles
- Insert [DATE] for unknown dates, [PARTY A] / [PARTY B] for parties, [JURISDICTION] for governing law
- Always separate sections with a line of dashes (----)

DO NOT invent regulatory requirements not present in the input context."""

# ── Data classes ──────────────────────────────────────────────────────────────

@dataclass
class AddendumItem:
    """Context for a single clause modification to be included in the addendum."""
    affected_name: str
    affected_ref: str
    suggestion: str
    original_clause_text: str | None
    change_summary: str
    article_ref: str
    change_type: str
    doc_title: str
    source: str
    publication_date: str | None
    contract_area: str | None


@dataclass
class AddendumContext:
    contract_name: str
    contract_area: str | None
    items: list[AddendumItem] = field(default_factory=list)

    @property
    def primary_doc_title(self) -> str:
        return self.items[0].doc_title if self.items else "Regulatory Change"

    @property
    def primary_source(self) -> str:
        return self.items[0].source if self.items else "Regulatory Authority"


# ── Public API ────────────────────────────────────────────────────────────────

async def fetch_context(alert_id: str, db: AsyncSession) -> list[AddendumContext]:
    """Load all contract-type ImpactItems for an alert and group by contract name.

    Returns a list of AddendumContext — one per unique contract in the alert.
    """
    rows = await _query_items(alert_id, db)

    groups: dict[str, AddendumContext] = {}
    for row in rows:
        name = row.affected_name
        if name not in groups:
            groups[name] = AddendumContext(
                contract_name=name,
                contract_area=row.contract_area,
            )
        pub_date = (
            row.publication_date.strftime("%B %d, %Y")
            if row.publication_date else None
        )
        groups[name].items.append(AddendumItem(
            affected_name=name,
            affected_ref=row.affected_ref or "—",
            suggestion=row.suggestion or "",
            original_clause_text=row.clause_text,
            change_summary=row.change_summary or "",
            article_ref=row.article_ref or "—",
            change_type=row.change_type or "other",
            doc_title=row.doc_title,
            source=row.source,
            publication_date=pub_date,
            contract_area=row.contract_area,
        ))

    return list(groups.values())


async def stream_draft(
    alert_id: str,
    db: AsyncSession,
) -> AsyncIterator[str]:
    """Stream a Claude-generated addendum as SSE events.

    Yields:
        data: {"type": "meta",    "contract": str, "doc_title": str}
        data: {"type": "token",   "text": str}
        data: {"type": "done"}
        data: {"type": "error",   "message": str}
    """
    import json

    contexts = await fetch_context(alert_id, db)
    if not contexts:
        yield f"data: {json.dumps({'type': 'error', 'message': 'No contract-type impact items found for this alert.'})}\n\n"
        return

    # Draft one addendum per contract (usually one, but may be multiple)
    for ctx in contexts:
        meta = {"type": "meta", "contract": ctx.contract_name, "doc_title": ctx.primary_doc_title}
        yield f"data: {json.dumps(meta)}\n\n"

        prompt = _build_prompt(ctx)

        if not settings.anthropic_api_key:
            fallback = _fallback_draft(ctx)
            for line in fallback.split("\n"):
                yield f"data: {json.dumps({'type': 'token', 'text': line + chr(10)})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            continue

        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
        try:
            with client.messages.stream(
                model=settings.llm_model,
                max_tokens=2048,
                system=_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": prompt}],
            ) as stream:
                for chunk in stream.text_stream:
                    yield f"data: {json.dumps({'type': 'token', 'text': chunk})}\n\n"
        except Exception as exc:
            logger.error("Claude streaming error in addendum draft: %s", exc)
            yield f"data: {json.dumps({'type': 'error', 'message': str(exc)})}\n\n"
            return

    yield f"data: {json.dumps({'type': 'done'})}\n\n"


def export_pdf(draft_text: str, title: str) -> bytes:
    """Render a draft text string as a formatted PDF document.

    Args:
        draft_text: The plain-text addendum (may contain ---- separators).
        title: Document title for the header.

    Returns:
        Raw PDF bytes.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        author="Regulatory Change Analyzer — AI Draft",
    )

    styles = getSampleStyleSheet()
    navy = colors.HexColor("#003366")
    gold = colors.HexColor("#C89B3C")

    title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=14, textColor=navy, spaceAfter=4)
    heading_style = ParagraphStyle("heading", parent=styles["Heading2"], fontSize=10, textColor=navy, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    body_style = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, leading=14, spaceAfter=6)
    disclaimer_style = ParagraphStyle("disc", parent=styles["Normal"], fontSize=7, textColor=colors.HexColor("#888888"), fontName="Helvetica-Oblique")

    story = []

    # Letterhead bar
    story.append(Paragraph("REGULATORY CHANGE ANALYZER", ParagraphStyle("lh", parent=styles["Normal"], fontSize=7, textColor=gold, fontName="Helvetica-Bold", letterSpacing=2)))
    story.append(Spacer(1, 0.15 * cm))
    story.append(HRFlowable(width="100%", thickness=1.5, color=navy))
    story.append(Spacer(1, 0.3 * cm))

    # Title
    story.append(Paragraph(title.upper(), title_style))
    story.append(Paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        ParagraphStyle("meta", parent=styles["Normal"], fontSize=8, textColor=colors.gray, spaceAfter=12),
    ))
    story.append(HRFlowable(width="100%", thickness=0.5, color=navy))
    story.append(Spacer(1, 0.4 * cm))

    # Body: split on ---- separators and article headings
    for block in draft_text.split("----"):
        block = block.strip()
        if not block:
            story.append(HRFlowable(width="100%", thickness=0.3, color=colors.HexColor("#d1d5db")))
            story.append(Spacer(1, 0.2 * cm))
            continue

        lines = block.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.15 * cm))
                continue

            # Detect section headings (ALL CAPS lines, ARTICLE / WHEREAS / NOW THEREFORE)
            upper = line.upper()
            is_heading = (
                line.startswith("ARTICLE")
                or line.startswith("WHEREAS")
                or line.startswith("NOW THEREFORE")
                or line.startswith("IN WITNESS WHEREOF")
                or (line == line.upper() and len(line) > 4 and len(line) < 80)
            )
            if is_heading:
                story.append(Paragraph(line, heading_style))
            else:
                # Escape XML special chars for ReportLab
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, body_style))

    # Footer disclaimer
    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#d1d5db")))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        "AI-GENERATED DRAFT — This document was produced by an automated language model and MUST be reviewed "
        "and approved by qualified legal counsel before execution. It does not constitute legal advice.",
        disclaimer_style,
    ))

    doc.build(story)
    logger.info("Addendum PDF generated (%d chars)", len(draft_text))
    return buf.getvalue()


def export_docx(draft_text: str, title: str) -> bytes:
    """Render a draft text string as a formatted DOCX document.

    Args:
        draft_text: The plain-text addendum (may contain ---- separators).
        title: Document title.

    Returns:
        Raw DOCX bytes.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    document = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = document.sections[0]
    section.page_width = int(21.0 * 914400 / 25.4)   # A4 width in EMUs
    section.page_height = int(29.7 * 914400 / 25.4)  # A4 height in EMUs
    section.left_margin = int(2.5 * 914400 / 25.4)
    section.right_margin = int(2.5 * 914400 / 25.4)
    section.top_margin = int(2.5 * 914400 / 25.4)
    section.bottom_margin = int(2.5 * 914400 / 25.4)

    navy = RGBColor(0, 51, 102)
    gold = RGBColor(200, 155, 60)

    # ── Header paragraph ─────────────────────────────────────────────────────
    hdr = document.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hdr.add_run("REGULATORY CHANGE ANALYZER")
    run.font.size = Pt(8)
    run.font.color.rgb = gold
    run.font.bold = True

    # ── Title ─────────────────────────────────────────────────────────────────
    p_title = document.add_heading(title.upper(), level=1)
    p_title.runs[0].font.color.rgb = navy

    # ── Date ──────────────────────────────────────────────────────────────────
    p_date = document.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    p_date.runs[0].font.size = Pt(8)
    p_date.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    document.add_paragraph()  # spacer

    # ── Body ──────────────────────────────────────────────────────────────────
    for block in draft_text.split("----"):
        block = block.strip()
        if not block:
            document.add_paragraph()
            continue

        lines = block.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                document.add_paragraph()
                continue

            is_heading = (
                line.startswith("ARTICLE")
                or line.startswith("WHEREAS")
                or line.startswith("NOW THEREFORE")
                or line.startswith("IN WITNESS WHEREOF")
                or (line == line.upper() and len(line) > 4 and len(line) < 80)
            )

            if is_heading:
                h = document.add_heading(line, level=2)
                for run in h.runs:
                    run.font.color.rgb = navy
            else:
                p = document.add_paragraph(line)
                p.runs[0].font.size = Pt(9.5) if p.runs else None

    # ── Disclaimer ────────────────────────────────────────────────────────────
    document.add_paragraph()
    disc = document.add_paragraph(
        "AI-GENERATED DRAFT — This document was produced by an automated language model and MUST be "
        "reviewed and approved by qualified legal counsel before execution. It does not constitute legal advice."
    )
    if disc.runs:
        disc.runs[0].font.size = Pt(7)
        disc.runs[0].font.color.rgb = RGBColor(140, 140, 140)
        disc.runs[0].font.italic = True

    buf = io.BytesIO()
    document.save(buf)
    logger.info("Addendum DOCX generated (%d chars)", len(draft_text))
    return buf.getvalue()


# ── Internal helpers ──────────────────────────────────────────────────────────

async def _query_items(alert_id: str, db: AsyncSession):
    """Fetch ImpactItems with full join context for an alert."""
    stmt = text(
        """
        SELECT
            ii.affected_name,
            ii.affected_ref,
            ii.suggestion,
            rc.summary       AS change_summary,
            rc.article_ref,
            rc.change_type,
            rc.new_text,
            rd.title         AS doc_title,
            rd.source,
            rd.publication_date,
            cc.text          AS clause_text,
            c.area           AS contract_area
        FROM impact_items ii
        JOIN impact_alerts ia  ON ia.id  = ii.alert_id
        JOIN regulatory_changes rc ON rc.id = ii.change_id
        JOIN regulatory_documents rd ON rd.id = rc.document_id
        LEFT JOIN contract_clauses cc ON cc.id = ii.clause_id
        LEFT JOIN contracts c ON c.name = ii.affected_name
        WHERE ia.id = CAST(:alert_id AS uuid)
          AND ii.impact_type = 'contract'
        ORDER BY ii.affected_name, ii.affected_ref
        """
    )
    result = await db.execute(stmt, {"alert_id": alert_id})
    return result.fetchall()


def _build_prompt(ctx: AddendumContext) -> str:
    """Format the addendum context into a structured prompt for Claude."""
    lines = [
        f"Draft a formal contract addendum for the following situation:",
        f"",
        f"CONTRACT: {ctx.contract_name}",
        f"AREA: {ctx.contract_area or 'Not specified'}",
        f"",
        "REGULATORY TRIGGER(S):",
    ]

    seen_docs: set[str] = set()
    for item in ctx.items:
        if item.doc_title not in seen_docs:
            lines.append(
                f"  - {item.source}: {item.doc_title}"
                + (f", {item.publication_date}" if item.publication_date else "")
                + (f", {item.article_ref}" if item.article_ref else "")
            )
            seen_docs.add(item.doc_title)
    lines.append("")
    lines.append("CLAUSE MODIFICATIONS REQUIRED:")

    for i, item in enumerate(ctx.items, 1):
        lines.append(f"")
        lines.append(f"  Modification {i}:")
        lines.append(f"  Clause: {item.affected_ref}")
        lines.append(f"  Regulatory change: {item.change_summary}")
        if item.original_clause_text:
            excerpt = item.original_clause_text[:400]
            lines.append(f"  Current clause text (excerpt): {excerpt}")
        lines.append(f"  Required action: {item.suggestion}")

    lines += [
        "",
        "Draft the complete addendum now, following the format in your instructions.",
        "Use '----' as a visual separator between major sections.",
    ]
    return "\n".join(lines)


def _fallback_draft(ctx: AddendumContext) -> str:
    """Minimal template draft when ANTHROPIC_API_KEY is not set."""
    items_text = "\n".join(
        f"ARTICLE {i} — AMENDMENT TO {item.affected_ref}\n"
        f"The following amendment is required: {item.suggestion}\n"
        for i, item in enumerate(ctx.items, 1)
    )
    return f"""ADDENDUM TO {ctx.contract_name.upper()}
[AI draft unavailable — set ANTHROPIC_API_KEY]

----

WHEREAS, {ctx.primary_source} has issued {ctx.primary_doc_title}, requiring amendments to the above contract;

NOW THEREFORE, the parties agree as follows:

----

{items_text}

----

GENERAL PROVISIONS
All other terms and conditions of the original agreement remain in full force and effect.

----

IN WITNESS WHEREOF, the parties have executed this Addendum as of [DATE].

PARTY A: ___________________________  Date: ___________

PARTY B: ___________________________  Date: ___________
"""
